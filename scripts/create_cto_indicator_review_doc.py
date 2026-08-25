from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "Indicator + Contracts" / "טיוטת Review ל-CTO - תנאים חוזיים ל-Indicator.docx"

NAVY = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(28, 39, 51)
MUTED = RGBColor(89, 96, 104)
LIGHT_BLUE = "E8EEF5"


def set_rtl_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = alignment
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_rtl_run(run):
    r_pr = run._element.get_or_add_rPr()
    rtl = r_pr.find(qn("w:rtl"))
    if rtl is None:
        rtl = OxmlElement("w:rtl")
        r_pr.append(rtl)
    rtl.set(qn("w:val"), "1")
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=160, bottom=100, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    table.autofit = False
    grid = table._tbl.tblGrid
    for grid_col in grid.gridCol_lst:
        grid_col.set(qn("w:w"), str(width_dxa))
    for row in table.rows:
        for cell in row.cells:
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width_dxa))
            tc_w.set(qn("w:type"), "dxa")


def set_spacing(paragraph, before=0, after=0, line=1.1):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(paragraph, text, *, bold=False, size=10.5, color=DARK):
    run = paragraph.add_run(text)
    set_rtl_run(run)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return run


def add_field_line(doc, label, value):
    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    set_spacing(p, after=1, line=1.05)
    add_text(p, f"{label}: ", bold=True, size=10.5, color=NAVY)
    add_text(p, value, size=10.5, color=DARK)
    return p


def add_condition(doc, clause, name, requirement, trigger, deadline, review_note):
    heading = doc.add_paragraph()
    set_rtl_paragraph(heading)
    set_spacing(heading, before=8, after=2, line=1.0)
    run = add_text(heading, f"סעיף {clause} - {name}", bold=True, size=12, color=NAVY)
    heading.paragraph_format.keep_with_next = True

    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    set_spacing(p, after=1, line=1.12)
    add_text(p, "דרישה חוזית: ", bold=True, size=10.5, color=DARK)
    add_text(p, requirement, size=10.5)

    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    set_spacing(p, after=1, line=1.12)
    add_text(p, "אירוע מפעיל: ", bold=True, size=10.5, color=DARK)
    add_text(p, trigger, size=10.5)
    add_text(p, "    |    דדליין: ", bold=True, size=10.5, color=DARK)
    add_text(p, deadline, size=10.5)

    p = doc.add_paragraph()
    set_rtl_paragraph(p)
    set_spacing(p, after=2, line=1.12)
    add_text(p, "נקודת Review: ", bold=True, size=10.5, color=MUTED)
    add_text(p, review_note, size=10.5, color=MUTED)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    set_rtl_paragraph(header)
    set_spacing(header, after=0, line=1.0)
    add_text(header, "BiDoc | Contracts → Indicator → Schedule", size=8.5, color=MUTED)

    footer = section.footer.paragraphs[0]
    set_rtl_paragraph(footer, WD_ALIGN_PARAGRAPH.CENTER)
    set_spacing(footer, after=0, line=1.0)
    add_text(footer, "טיוטה פנימית - לדיון ואישור בלבד", size=8.5, color=MUTED)

    title = doc.add_paragraph()
    set_rtl_paragraph(title)
    set_spacing(title, after=3, line=1.0)
    add_text(title, "טיוטת Review ל-CTO", bold=True, size=22, color=NAVY)

    subtitle = doc.add_paragraph()
    set_rtl_paragraph(subtitle)
    set_spacing(subtitle, after=10, line=1.0)
    add_text(subtitle, "תנאים חוזיים מוצעים להעברה ל-Indicator", size=13, color=MUTED)

    add_field_line(doc, "מטרה", "לאשר אילו התחייבויות חוזיות יעברו ל-Indicator לצורך זיהוי trigger ומעקב.")
    add_field_line(doc, "סטטוס", "טיוטת הצעות בלבד - טרם נכתבו החלטות או רשומות לטבלת היעד.")

    callout = doc.add_table(rows=1, cols=1)
    set_table_width(callout)
    cell = callout.cell(0, 0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    set_rtl_paragraph(p)
    set_spacing(p, after=0, line=1.12)
    add_text(p, "המלצה: ", bold=True, size=10.5, color=NAVY)
    add_text(p, "להעביר שמונה תנאים ל-Indicator לאחר אישור; להשאיר את סעיף 13.4 ב-Contracts בלבד.", size=10.5, color=DARK)

    section_title = doc.add_paragraph()
    set_rtl_paragraph(section_title)
    set_spacing(section_title, before=12, after=2, line=1.0)
    add_text(section_title, "מועמדים להעברה ל-Indicator", bold=True, size=15, color=BLUE)

    candidates = [
        ("3.6", "הכנת לוח זמנים מפורט", "הקבלן מגיש לוח זמנים מפורט.", "חתימת החוזה", "10 ימים קלנדריים", "לוודא שהחתימה זמינה כאירוע מתועד למנוע."),
        ("3.8", "תגבור צוות עובדים", "הקבלן מתגבר את הצוות לפי דרישת המזמין.", "דרישה מפורשת של המזמין", "2 ימי עבודה", "לזהות דרישה אופרטיבית ולא אזכור כללי של כוח אדם; תלוי בלוח חגים."),
        ("8.2", "תקופת בדק", "אחריות הקבלן לטיב העבודות והחומרים.", "סיום כלל העבודות", "12 חודשים", "להגדיר אם אירוע הסיום הקנוני הוא סיום, מסירה או אישור מפקח."),
        ("8.5", "החזר הוצאות בגין ליקויים", "הקבלן מחזיר למזמין הוצאות תיקון ליקויים.", "דרישה ראשונה של המזמין", "7 ימים קלנדריים", "ה-trigger הוא דרישת תשלום או החזר מתועדת, ולא גילוי הליקוי."),
        ("8.10.2.1", "תיקון פגם במערכות תומכות", "תיקון פגם במערכות תומכות.", "קבלת הודעת המפקח", "12 שעות", "SLA קצר; יש לזהות מועד קבלה ולא רק מועד יצירת ההודעה."),
        ("8.10.2.1", "תיקון פגם במערכות רלוונטיות", "תיקון פגם במערכות רלוונטיות.", "קבלת הודעת המפקח", "7 ימים קלנדריים", "תנאי נפרד מאותו סעיף משום שמשך הזמן שונה."),
        ("8.10.2.4", "בדיקה נוספת של איכות העבודות", "בדיקה נוספת של טיב העבודות.", "מסירת תעודת השלמה לפרויקט או לחלק ממנו", "24 חודשים", "לאשר אם תעודה לחלק מהפרויקט יוצרת תנאי נפרד לכל חלק."),
        ("15.3", "חזקת מסירה בדואר רשום", "הודעה בדואר רשום נחשבת כאילו נמסרה.", "תאריך משלוח הדואר הרשום", "5 ימים קלנדריים", "לוודא שמקור ה-trigger מספק תאריך משלוח מהימן."),
    ]
    for candidate in candidates:
        add_condition(doc, *candidate)

    excluded_heading = doc.add_paragraph()
    set_rtl_paragraph(excluded_heading)
    set_spacing(excluded_heading, before=12, after=2, line=1.0)
    add_text(excluded_heading, "מועמד שלא יועבר ל-Indicator", bold=True, size=15, color=BLUE)
    add_condition(
        doc,
        "13.4",
        "איסור פנייה ללקוחות",
        "איסור פנייה ללקוחות המזמינה לאחר סיום ההתקשרות.",
        "סיום תקופת ההתקשרות",
        "6 חודשים",
        "לא מועבר ל-Indicator: זו התחייבות מסחרית או משפטית, ולא תנאי תפעולי בלוח הזמנים.",
    )

    final_heading = doc.add_paragraph()
    set_rtl_paragraph(final_heading)
    set_spacing(final_heading, before=12, after=2, line=1.0)
    add_text(final_heading, "החלטה נדרשת מה-CTO", bold=True, size=15, color=BLUE)
    final = doc.add_paragraph()
    set_rtl_paragraph(final)
    set_spacing(final, after=0, line=1.15)
    add_text(final, "לאשר את העוגן, משך הזמן והרלוונטיות התפעולית של שמונת המועמדים. לאחר אישור, ייווצרו רשומות pending בטבלת היעד; ה-Indicator יזהה את ה-trigger ויחשב את התאריך בפועל.", size=10.5, color=DARK)

    doc.core_properties.title = "טיוטת Review ל-CTO - תנאים חוזיים ל-Indicator"
    doc.core_properties.subject = "Contracts → Indicator → Schedule"
    doc.core_properties.author = "BiDoc"
    doc.save(OUT)
    print("Document created successfully")


if __name__ == "__main__":
    main()
