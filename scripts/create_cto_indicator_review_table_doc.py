from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "Indicator + Contracts" / "טיוטת Review ל-CTO - תנאים חוזיים ל-Indicator (טבלה).docx"

NAVY = RGBColor(31, 77, 120)
DARK = RGBColor(28, 39, 51)
MUTED = RGBColor(89, 96, 104)
HEADER_FILL = "E8EEF5"


def set_rtl_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = alignment
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def add_run(paragraph, text, *, bold=False, size=9.5, color=DARK):
    run = paragraph.add_run(text)
    r_pr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    r_pr.append(rtl)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    return run


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_geometry(table, widths):
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    table.autofit = False
    for grid_col, width in zip(table._tbl.tblGrid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_cell_text(cell, text, *, bold=False, color=DARK, size=9.5):
    p = cell.paragraphs[0]
    set_rtl_paragraph(p)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    add_run(p, text, bold=bold, size=size, color=color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Arial")
    normal.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    set_rtl_paragraph(header)
    add_run(header, "BiDoc | Contracts → Indicator → Schedule", size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    set_rtl_paragraph(footer, WD_ALIGN_PARAGRAPH.CENTER)
    add_run(footer, "טיוטה פנימית - לדיון ואישור בלבד", size=8.5, color=MUTED)

    title = doc.add_paragraph()
    set_rtl_paragraph(title)
    title.paragraph_format.space_after = Pt(2)
    add_run(title, "טיוטת Review ל-CTO", bold=True, size=21, color=NAVY)
    subtitle = doc.add_paragraph()
    set_rtl_paragraph(subtitle)
    subtitle.paragraph_format.space_after = Pt(8)
    add_run(subtitle, "תנאים חוזיים מוצעים להעברה ל-Indicator", size=12.5, color=MUTED)
    summary = doc.add_paragraph()
    set_rtl_paragraph(summary)
    summary.paragraph_format.space_after = Pt(10)
    add_run(summary, "החלטה מוצעת: ", bold=True, size=10.5, color=NAVY)
    add_run(summary, "להעביר שמונה תנאים ל-Indicator לאחר אישור; להשאיר את סעיף 13.4 ב-Contracts בלבד.", size=10.5)

    rows = [
        ("3.6", "הכנת לוח זמנים מפורט", "חתימת החוזה\n10 ימים קלנדריים", "להעביר ל-Indicator.\nלוודא שהחתימה מתועדת."),
        ("3.8", "תגבור צוות עובדים לפי דרישת המזמין", "דרישה מפורשת של המזמין\n2 ימי עבודה", "להעביר ל-Indicator.\nנדרש לוח חגים וזיהוי דרישה אופרטיבית."),
        ("8.2", "תקופת בדק לטיב העבודות והחומרים", "סיום כלל העבודות\n12 חודשים", "להעביר ל-Indicator.\nלקבע אירוע סיום קנוני."),
        ("8.5", "החזר הוצאות למזמין בגין תיקון ליקויים", "דרישה ראשונה של המזמין\n7 ימים קלנדריים", "להעביר ל-Indicator.\nה-trigger הוא דרישה מתועדת."),
        ("8.10.2.1", "תיקון פגם במערכות תומכות", "קבלת הודעת המפקח\n12 שעות", "להעביר ל-Indicator.\nSLA קצר: לזהות מועד קבלה."),
        ("8.10.2.1", "תיקון פגם במערכות רלוונטיות", "קבלת הודעת המפקח\n7 ימים קלנדריים", "להעביר ל-Indicator.\nרשומה נפרדת משום שהמשך שונה."),
        ("8.10.2.4", "בדיקה נוספת של טיב העבודות", "מסירת תעודת השלמה\n24 חודשים", "להעביר ל-Indicator.\nלאשר אם חלק פרויקט יוצר תנאי נפרד."),
        ("13.4", "איסור פנייה ללקוחות המזמינה", "סיום ההתקשרות\n6 חודשים", "לא להעביר.\nהתחייבות מסחרית/משפטית, לא תפעולית."),
        ("15.3", "חזקת מסירה בדואר רשום", "תאריך המשלוח\n5 ימים קלנדריים", "להעביר ל-Indicator.\nנדרש מקור משלוח מהימן."),
    ]
    table = doc.add_table(rows=1, cols=4)
    widths = [900, 2500, 2500, 3460]
    set_table_geometry(table, widths)
    headers = ["סעיף", "תנאי חוזי", "אירוע מפעיל ודדליין", "החלטה מוצעת ונקודת Review"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_shading(cell, HEADER_FILL)
        add_cell_text(cell, text, bold=True, color=NAVY, size=9.5)
    set_repeat_header(table.rows[0])
    for row_values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_values):
            add_cell_text(cell, text)
        if row_values[0] == "13.4":
            for cell in cells:
                set_shading(cell, "F7F3F3")
    set_table_geometry(table, widths)

    note = doc.add_paragraph()
    set_rtl_paragraph(note)
    note.paragraph_format.space_before = Pt(10)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.line_spacing = 1.12
    add_run(note, "לאחר אישור: ", bold=True, size=10.5, color=NAVY)
    add_run(note, "לכל תנאי מאושר תיווצר רשומת pending בטבלת היעד. ה-Indicator נשאר אחראי לזיהוי האירוע ולחישוב התאריך בפועל.", size=10.5)

    doc.core_properties.title = "טיוטת Review ל-CTO - תנאים חוזיים ל-Indicator (טבלה)"
    doc.core_properties.subject = "Contracts → Indicator → Schedule"
    doc.core_properties.author = "BiDoc"
    doc.save(OUT)
    print("Table document created successfully")


if __name__ == "__main__":
    main()
